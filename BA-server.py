from mcp.server.fastmcp import FastMCP
import json
import os
import sys
import uuid
import datetime
from dotenv import load_dotenv
from langchain_huggingface import HuggingFaceEmbeddings
from pinecone import Pinecone
from tavily import TavilyClient
from PyPDF2 import PdfReader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Load environment variables
load_dotenv()

# Validate required environment variables
required_env_vars = ['PINECONE_API_KEY', 'TAVILY_API_KEY']
missing_vars = [var for var in required_env_vars if not os.environ.get(var)]
if missing_vars:
    print(f"Error: Missing required environment variables: {', '.join(missing_vars)}", file=sys.stderr)
    sys.exit(1)

# Connection to Pinecone
pinecone_api_key = os.environ.get('PINECONE_API_KEY')
pc = Pinecone(api_key=pinecone_api_key)
index = pc.Index("test1")
embedder = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")

# Directory setup
FILESYSTEM_PATH = None

text_splitter = RecursiveCharacterTextSplitter(
    chunk_size=1000,
    chunk_overlap=200,
    separators=["\n\n", ".", " ", ""]
)


# Utility functions
def sanitize_filename(filename: str) -> str:
    """Sanitize filename to prevent path traversal attacks."""
    # Remove path separators and other potentially dangerous characters
    import re
    sanitized = re.sub(r'[<>:"/\\|?*]', '_', filename)
    # Remove leading/trailing dots and spaces
    sanitized = sanitized.strip('. ')
    # Ensure filename is not empty
    return sanitized if sanitized else 'unnamed'



def extract_text_from_pdf(pdf_path: str) -> str:
    reader = PdfReader(pdf_path)
    text = ""
    for page in reader.pages:
        if page.extract_text():
            text += page.extract_text()
    return text

def extract_text_from_pdf(pdf_path: str) -> str:
    reader = PdfReader(pdf_path)
    return "".join(page.extract_text() or "" for page in reader.pages)

# Create MCP server
mcp = FastMCP("Broadaxis-Server" , host="0.0.0.0",  port = 8892)

@mcp.tool()
def set_filesystem_path(path: str) -> str:
    """
    Sets the base filesystem path for file tools (e.g., PDF generation, RAG ingestion).
    Should be called once from Claude Desktop with the local path the assistant can access.
    """
    import os
    global FILESYSTEM_PATH

    global FILESYSTEM_PATH
    FILESYSTEM_PATH = path
    return json.dumps({
        "status": "success",
        "message": f"Filesystem path set to: {FILESYSTEM_PATH}"
    })


@mcp.tool()
def get_current_filesystem_path() -> str:
    return json.dumps({
        "filesystem_path": FILESYSTEM_PATH or "Not set"
    })


@mcp.tool()
def create_custom_rag_from_local_pdfs() -> str:
    """
    Scans the local Claude-accessible folder for PDF files,
    extracts content, chunks, embeds, and upserts to Pinecone.
    """
    if not FILESYSTEM_PATH:
        return json.dumps({"status": "error", "message": "Filesystem path not set."})
    try:
        processed = []
        for filename in os.listdir(FILESYSTEM_PATH):
            if filename.lower().endswith(".pdf"):
                path = os.path.join(FILESYSTEM_PATH, filename)
                text = extract_text_from_pdf(path)
                chunks = text_splitter.split_text(text)
                for i, chunk in enumerate(chunks):
                    embedding = embedder.embed_documents([chunk])[0]
                    metadata = {"source": filename, "text": chunk}
                    unique_id = f"{filename}_chunk_{i}"
                    index.upsert([(unique_id, embedding, metadata)])
                processed.append(filename)
        return json.dumps({
            "status": "success",
            "processed_pdfs": processed,
            "message": f"Indexed {len(processed)} PDFs."
        })
    except Exception as e:
        return json.dumps({"status": "error", "error": str(e)})


@mcp.tool()
def Broadaxis_knowledge_search(query: str):
    """
    Retrieves the most relevant company's (Broadaxis) information from the internal knowledge base in response to a company related query.
    This tool performs semantic search over a RAG-powered database containing details about Broadaxis's background, team, projects, responsibilities, and domain expertise. It is designed to support tasks such as retrieving the knowledge regarding the company, surfacing domain-specific experience.

    Args:
        query: A natural language request related to the company’s past work, expertise, or capabilities (e.g., "What are the team's responsibilities?").
    """
    try:
        query_embedding = embedder.embed_documents([query])[0]
        query_result = index.query(
            vector=[query_embedding],
            top_k=5,
            include_metadata=True,
            namespace=""
        )
        documents = [result['metadata']['text'] for result in query_result['matches']]
        return documents
    except Exception as e:
        return json.dumps({"error": f"Knowledge search failed: {str(e)}"})    


tavily = TavilyClient(api_key=os.environ["TAVILY_API_KEY"]) 

@mcp.tool()
def web_search_tool(query: str):
    """
    Performs a real-time web search using Tavily and returns relevant results
    (including title, URL, and snippet).

    Args:
        query: A natural language search query.

    Returns:
        A JSON string with the top search results.
    """
    try:
        # Perform the search
        results = tavily.search(query=query, search_depth="advanced", include_answer=False)

        # Extract and format results
        formatted = [
            {
                "title": r.get("title"),
                "url": r.get("url"),
                "snippet": r.get("content")
            }
            for r in results.get("results", [])
        ]

        return json.dumps({"results": formatted})

    except Exception as e:
        return json.dumps({"error": str(e)})
    
@mcp.tool()
def generate_pdf_document(title: str, content: str, filename: str = None) -> str:
    """Generate a PDF document and save it in Claude's filesystem folder."""
    if not FILESYSTEM_PATH:
        return json.dumps({"status": "error", "message": "Filesystem path not set."})
    try:
        filename = sanitize_filename((filename or f"document_{uuid.uuid4().hex[:8]}").replace('.pdf', ''))
        file_path = os.path.join(FILESYSTEM_PATH, f"{filename}.pdf")
        # file_path = os.path.normpath(os.path.join(FILESYSTEM_PATH, f"{filename}.pdf"))
        os.makedirs(os.path.dirname(file_path), exist_ok=True)

        if not os.access(os.path.dirname(file_path), os.W_OK):
            return json.dumps({"status": "error", "error": f"No write permission to {os.path.dirname(file_path)}"})
        doc = SimpleDocTemplate(file_path, pagesize=letter)
        styles = getSampleStyleSheet()
        story = [Paragraph(title, ParagraphStyle('Title', parent=styles['Heading1'], fontSize=18, alignment=1))]
        story.append(Spacer(1, 20))
        for line in content.split('\n'):
            if line.strip():
                if line.startswith('# '):
                    story.append(Paragraph(line[2:], styles['Heading1']))
                elif line.startswith('## '):
                    story.append(Paragraph(line[3:], styles['Heading2']))
                elif line.startswith('### '):
                    story.append(Paragraph(line[4:], styles['Heading3']))
                elif line.startswith(('- ', '* ')):
                    story.append(Paragraph("• " + line[2:], styles['Normal']))
                else:
                    story.append(Paragraph(line, styles['Normal']))
                story.append(Spacer(1, 6))
        doc.build(story)
        return json.dumps({
            "status": "success",
            "filename": f"{filename}.pdf",
            "file_path": file_path,
            "file_size": os.path.getsize(file_path),
            "download_url": f"/download/{filename}.pdf",
            "created_at": datetime.datetime.now().isoformat(),
            "type": "pdf"
        })
    except Exception as e:
        return json.dumps({"status": "error", "error": str(e)})

@mcp.tool()
def generate_word_document(title: str, content: str, filename: str = None) -> str:
    """Generate a Word document and save it in Claude's filesystem folder."""
    if not FILESYSTEM_PATH:
        return json.dumps({"status": "error", "message": "Filesystem path not set."})
    try:
        filename = sanitize_filename((filename or f"document_{uuid.uuid4().hex[:8]}").replace('.docx', ''))
        file_path = os.path.join(FILESYSTEM_PATH, f"{filename}.docx")
        # file_path = os.path.normpath(os.path.join(FILESYSTEM_PATH, f"{filename}.docx"))
        os.makedirs(os.path.dirname(file_path), exist_ok=True)

        if not os.access(os.path.dirname(file_path), os.W_OK):
            return json.dumps({"status": "error", "error": f"No write permission to {os.path.dirname(file_path)}"})
        doc = Document()
        title_paragraph = doc.add_heading(title, level=1)
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()
        for line in content.split('\n'):
            line = line.strip()
            if line:
                if line.startswith('# '):
                    doc.add_heading(line[2:], level=1)
                elif line.startswith('## '):
                    doc.add_heading(line[3:], level=2)
                elif line.startswith('### '):
                    doc.add_heading(line[4:], level=3)
                elif line.startswith(('- ', '* ')):
                    doc.add_paragraph(line[2:], style='List Bullet')
                elif any(line.startswith(f"{i}. ") for i in range(1, 100)):
                    doc.add_paragraph(line[line.find('. ') + 2:], style='List Number')
                else:
                    doc.add_paragraph(line)
            else:
                doc.add_paragraph()
        doc.save(file_path)
        return json.dumps({
            "status": "success",
            "filename": f"{filename}.docx",
            "file_path": file_path,
            "file_size": os.path.getsize(file_path),
            "download_url": f"/download/{filename}.docx",
            "created_at": datetime.datetime.now().isoformat(),
            "type": "docx"
        })
    except Exception as e:
        return json.dumps({"status": "error", "error": str(e)})



@mcp.tool()
def generate_text_file(content: str, filename: str = None, file_extension: str = "txt") -> str:
    """Generate a text-based file (txt, md, csv, json, etc.) and save it in Claude's folder."""
    if not FILESYSTEM_PATH:
        return json.dumps({"status": "error", "message": "Filesystem path not set."})
    try:
        filename = sanitize_filename((filename or f"file_{uuid.uuid4().hex[:8]}").split('.')[0])
        file_path = os.path.join(FILESYSTEM_PATH, f"{filename}.{file_extension}")
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(content)
        return json.dumps({
            "status": "success",
            "filename": f"{filename}.{file_extension}",
            "file_path": file_path,
            "file_size": os.path.getsize(file_path),
            "download_url": f"/download/{filename}.{file_extension}",
            "created_at": datetime.datetime.now().isoformat(),
            "type": file_extension
        })
    except Exception as e:
        return json.dumps({"status": "error", "error": str(e)})




@mcp.prompt(title="Step 1: Identifying the Documents")
def summarize_documents():
    """read PDFs from filesystem path, categorize them as RFP/RFI/RFQ-related, fillable forms, or non-fillable documents."""
    return f"""
You are BroadAxis-AI, an intelligent assistant designed to analyze procurement documents and assist with document classification.

As the **first step**, review the files from the provided `filesystem_path` and use PDF processing tools to categorize each uploaded PDF into the following groups:

1. 📘 **Primary Documents** — PDFs that contain RFP, RFQ, or RFI content (e.g., project scope, requirements, evaluation criteria).
2. 📝 **Fillable Forms** — PDFs with interactive fields intended for user input (e.g., pricing tables, response forms).
3. 📄 **Non-Fillable Documents** — PDFs that are neither RFP-type nor interactive, such as attachments or informational appendices.

---

Once the classification is complete:

📊 **Would you like to proceed to the next step and generate summaries for the relevant documents?**  
If yes, please upload the files and attach the summary prompt template.
"""



@mcp.prompt(title="Step-2 : Summarize Uploaded Document(s)")
def summarize_documents():
    """Generate a structured summary of uploaded RFP, RFQ, or RFI documents."""
    return f"""You are BroadAxis-AI, a professional assistant trained to analyze and summarize procurement documents such as RFPs (Request for Proposals), RFQs (Request for Quotations), and RFIs (Request for Information).
Your task is to review the uploaded document(s) and provide a **structured, factual summary** using only the information found in each file. **Do not infer, assume, or hallucinate any content.**
For each document, present the following sections clearly and professionally: 
if you dont find the sections below, just include the summary with key details you find, and generate a general summary, the whole point is for the user to understand the what the document is abut. 
---
### 📄 Document [Auto Numbered]

**1. Introduction**  
Briefly describe what this document is about.

**2. Document Title**  
(Extract the title of the document, if available.)

**3. Summary**  
Provide a concise overview of the document's purpose, objectives, and key highlights.

**4. Key Requirements**  
- Present a bulleted list of technical, functional, or compliance requirements.
- Use direct quotes or paraphrased bullet points from the document.

**5. Project Scope & Timeline**  
Summarize the scope of work and any milestones, phases, or deadlines provided.

**6. Budget Information**  
State whether budget, pricing expectations, or cost caps are mentioned.  
*If not available, write:* “Not specified in the document.”

**7. Evaluation Criteria**  
Detail how proposals will be assessed — e.g., by cost, technical merit, experience, etc.  
*If not available, write:* “Not specified in the document.”

**8. Submission Deadline**  
Include the exact date and time if available.  
*If not available, write:* “Not specified in the document.”

**9. Key Risks & Opportunities**  
Note any mentioned risks, constraints, or unique opportunities outlined in the document.  
*If not available, write:* “Not specified in the document.”

**10. Contact Information**  
Include name(s), email, phone number, or other official contact details provided.  
*If not available, write:* “Not specified in the document.”

---

Once all documents have been summarized, conclude with the following question to the user:

📊 **Would you like additional insights or a Go/No-Go recommendation for any of these opportunities?**
"""

@mcp.prompt(title="Step-3 : Go/No-Go Recommendation")
def go_no_go_recommendation() -> str:
    return """
You are BroadAxis-AI, an assistant trained to evaluate whether BroadAxis should pursue an RFP, RFQ, or RFI opportunity.
The user has uploaded one or more opportunity documents. You have already summarized them/if not ask for the user to upload RFP/RFI/RF documents and generate summary.
Now perform a structured **Go/No-Go analysis** using the following steps:
---
### 🧠 Step-by-Step Evaluation Framework

1. **Review the RFP Requirements**
   - Highlight the most critical needs and evaluation criteria.

2. **Search Internal Knowledge** (via Broadaxis_knowledge_search)
   - Identify relevant past projects
   - Retrieve proof of experience in similar domains
   - Surface known strengths or capability gaps

3. **Evaluate Capability Alignment**
   - Estimate percentage match (e.g., "BroadAxis meets ~85% of the requirements")
   - Note any missing capabilities or unclear requirements

4. **Assess Resource Requirements**
   - Are there any specialized skills, timelines, or staffing needs?
   - Does BroadAxis have the necessary team or partners?

5. **Evaluate Competitive Positioning**
   - Based on known experience and domain, would BroadAxis be competitive?

Use only verified internal information (via Broadaxis_knowledge_search) and the uploaded documents.
Do not guess or hallucinate capabilities. If information is missing, clearly state what else is needed for a confident decision.
if your recommendation is a Go, list down the things to the user of the tasks he need to complete  to finish the submission of RFP/RFI/RFQ. 

"""

@mcp.prompt(title="Step-4 : Generate Proposal or Capability Statement")
def generate_capability_statement() -> str:
    return """
You are BroadAxis-AI, an assistant trained to generate high-quality capability statements and proposal documents for RFP and RFQ responses.
The user has either uploaded an opportunity document or requested a formal proposal. Use all available information from:

- Uploaded documents (RFP/RFQ)
- Internal knowledge (via Broadaxis_knowledge_search)
- Prior summaries or analyses already provided

---

### 🧠 Instructions

- Do not invent names, projects, or facts.
- Use Broadaxis_knowledge_search to populate all relevant content.
- Leave placeholders where personal or business info is not available.
- Maintain professional, confident, and compliant tone.

If this proposal is meant to be saved, offer to generate a PDF or Word version using the appropriate tool.

"""

@mcp.prompt(title="Step-5 : Fill in Missing Information")
def fill_missing_information() -> str:
    return """
You are BroadAxis-AI, an intelligent assistant designed to fill in missing fields using ppdf filler tool , answer RFP/RFQ questions, and complete response templates **strictly using verified information**.
 Your task is to **complete the missing sections** on the fillable docuemnst which you have identified previously with reliable information from:

1. Broadaxis_knowledge_search (internal database)
2. The uploaded document itself
3. Prior chat context (if available)
---
### 🧠 RULES (Strict Compliance)

- ❌ **DO NOT invent or hallucinate** company details, financials, certifications, team names, or client info.
- ❌ **DO NOT guess** values you cannot verify.
- 🔐 If the question involves **personal, legal, or confidential information**, **do not fill it**.
- ✅ Use internal knowledge only when it clearly answers the field.
---
### ✅ Final Instruction
Only fill what you can verify using Broadaxis_knowledge_search and uploaded content. Leave everything else with clear, professional placeholders.

"""

@mcp.prompt(title="📁 Set Local Filesystem Path")
def set_local_filesystem_path_prompt(path: str = "C:/Users/rohka/OneDrive/Desktop/test-mcp") -> str:
    return f"""
You are setting the local folder path that BroadAxis-AI will use to read and write files.

Set the following path as your working directory:
`{path}`
"""

@mcp.prompt(title="📂 Show Current Filesystem Path")
def show_current_filesystem_path_prompt() -> str:
    return """
You're checking the currently configured local folder path that BroadAxis-AI uses.

Trigger the `get_current_filesystem_path()` tool to see the folder currently being used to store and retrieve files (PDFs, DOCXs, etc).
"""

@mcp.prompt(title="🔍 Create Custom RAG from Local PDF")
def create_rag_from_pdf_prompt(pdf_file: str = "example.pdf") -> str:
    return f"""
You are creating a custom RAG dataset by extracting text from a local PDF file.

Please ensure the PDF file named:
"""

@mcp.prompt(title="Configure Word Document Formatting")
def configure_word_formatting(
    font: str = "Calibri",
    font_size: int = 11,
    heading_font_size: int = 14,
    heading_bold: bool = True,
    line_spacing: float = 1.15,
    margin_top: str = "1 inch",
    margin_bottom: str = "1 inch",
    margin_left: str = "1 inch",
    margin_right: str = "1 inch",
    page_orientation: str = "portrait"
) -> str:
    """
    Prompt to configure Word document formatting settings.
    Users can override any of the default values.
    """
    return f"""Please use the following Word document formatting settings:

### 📝 Formatting Instructions

- **Font**: {font}
- **Font Size**: {font_size} pt
- **Heading Font Size**: {heading_font_size} pt
- **Heading Bold**: {"Yes" if heading_bold else "No"}
- **Line Spacing**: {line_spacing}
- **Margins**:
  - Top: {margin_top}
  - Bottom: {margin_bottom}
  - Left: {margin_left}
  - Right: {margin_right}
- **Page Orientation**: {page_orientation.capitalize()}

Apply these settings when generating Word documents, including proposals, summaries, and responses. Ensure consistency and professionalism throughout the layout.
"""


if __name__ == "__main__":
    import logging
    logging.basicConfig(level=logging.DEBUG)
    mcp.run(transport="sse")